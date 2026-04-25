"""
rfp_text_pipeline.py — Daily pull of RFP attachment text for 541511/541512 opportunities.

Bulk-search mode. One API call returns up to 1,000 opportunities with their
resourceLinks attached — so the full daily catch typically needs just a handful
of calls, not one-per-opp.

Flow:
  1. Pull state from R2 (if configured): processed.json, last_fetched_date.json.
  2. Determine the posted-date window: from (last_fetched_date - 1 day) to today.
     First run falls back to --start-date or 180 days ago.
  3. Paginate through api.sam.gov/prod/opportunities/v2/search with
     postedFrom/postedTo/limit=1000/offset= until the page is short.
  4. For each opp:
       - Skip if NAICS prefix doesn't match (default: 541511, 541512).
       - Skip if notice type isn't in the RFP-adjacent set.
       - Skip if noticeId already in processed.
       - Download each resourceLink (free S3), extract text via pypdf /
         python-docx / openpyxl. Run the regex label classifier.
       - Write bundles/{noticeId}.json and add to processed.
  5. Update last_fetched_date = today. Push state + bundles to R2.

Run:
    python3 rfp_text_pipeline.py                       # default daily window
    python3 rfp_text_pipeline.py --start-date 2025-10-01   # bootstrap
    python3 rfp_text_pipeline.py --dry-run             # no API, no downloads
    python3 rfp_text_pipeline.py --max-api-calls 5     # cap paginator
"""

from __future__ import annotations

import argparse
import csv
import hashlib
import io
import json
import os
import re
import sys
import time
from datetime import date, datetime, timedelta, timezone
from pathlib import Path
from typing import Any, Iterator

import requests
from dotenv import load_dotenv

try:
    from pypdf import PdfReader
except ImportError as exc:  # pragma: no cover
    raise SystemExit("pypdf missing — pip install pypdf") from exc

try:
    from docx import Document  # python-docx
except ImportError:
    Document = None

try:
    from openpyxl import load_workbook  # xlsx
except ImportError:
    load_workbook = None

load_dotenv()

csv.field_size_limit(min(sys.maxsize, 2**31 - 1))


# ---------------------------------------------------------------------------
# Config
# ---------------------------------------------------------------------------

DATA_DIR       = Path("data/rfp_text")
STATE_DIR      = DATA_DIR / "state"
BUNDLE_DIR     = DATA_DIR / "bundles"
PROCESSED_JSON   = STATE_DIR / "processed.json"
LAST_DATE_JSON   = STATE_DIR / "last_fetched_date.json"
QUOTA_JSON       = STATE_DIR / "quota.json"
SCAN_CURSOR_JSON = STATE_DIR / "scan_cursor.json"  # pinned window + offset for multi-day drains

OPP_SEARCH_API = "https://api.sam.gov/prod/opportunities/v2/search"

DEFAULT_NOTICE_TYPES = {
    "Solicitation",
    "Combined Synopsis/Solicitation",
    "Sources Sought",
    "Presolicitation",
    "Special Notice",
    # Justification + Fair Opportunity / Limited Sources Justification live in
    # the pull_usaspending repo instead — they're about non-competition
    # rationale on awarded contracts, not RFP language.
}
DEFAULT_NAICS_PREFIXES = ("541511", "541512")
DEFAULT_LOOKBACK_DAYS  = 180

API_KEY = os.environ.get("SAM_API_KEY")
R2_PREFIX = "it_rfps/"


# ---------------------------------------------------------------------------
# State I/O
# ---------------------------------------------------------------------------

def _load_json(path: Path, default: Any) -> Any:
    if path.exists() and path.stat().st_size > 0:
        return json.loads(path.read_text())
    return default


def _save_json(path: Path, data: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(data, indent=2, default=str))


def _maybe_hydrate_from_r2() -> None:
    if not os.environ.get("CF_R2_ACCOUNT_ID"):
        return
    if PROCESSED_JSON.exists() and LAST_DATE_JSON.exists():
        return
    print("Hydrating state from R2...")
    import r2_sync

    s3 = r2_sync._client()
    paginator = s3.get_paginator("list_objects_v2")
    for page in paginator.paginate(Bucket=r2_sync.BUCKET, Prefix=R2_PREFIX + "state/"):
        for obj in page.get("Contents", []):
            name = Path(obj["Key"]).name
            local = STATE_DIR / name
            local.parent.mkdir(parents=True, exist_ok=True)
            s3.download_file(r2_sync.BUCKET, obj["Key"], str(local))
            print(f"  R2 -> {name}")


def _clear_scan_cursor() -> None:
    """Remove scan_cursor state locally + on R2 (called after a full-window drain)."""
    if SCAN_CURSOR_JSON.exists():
        SCAN_CURSOR_JSON.unlink()
    if not os.environ.get("CF_R2_ACCOUNT_ID"):
        return
    import r2_sync
    key = R2_PREFIX + "state/" + SCAN_CURSOR_JSON.name
    try:
        r2_sync._client().delete_object(Bucket=r2_sync.BUCKET, Key=key)
        print(f"  scan_cursor cleared on R2")
    except Exception as e:
        print(f"  warning: could not delete R2 scan_cursor ({e})")


def _push_state_to_r2() -> None:
    if not os.environ.get("CF_R2_ACCOUNT_ID"):
        return
    import r2_sync

    s3 = r2_sync._client()
    for f in STATE_DIR.glob("*.json"):
        key = R2_PREFIX + "state/" + f.name
        s3.upload_file(str(f), r2_sync.BUCKET, key)
    print(f"  state -> R2")


def _push_bundles_to_r2(noticeids: list[str]) -> None:
    if not os.environ.get("CF_R2_ACCOUNT_ID") or not noticeids:
        return
    import r2_sync

    s3 = r2_sync._client()
    for nid in noticeids:
        local = BUNDLE_DIR / f"{nid}.json"
        if not local.exists():
            continue
        key = R2_PREFIX + f"bundles/{nid}.json"
        s3.upload_file(str(local), r2_sync.BUCKET, key)
    print(f"  {len(noticeids)} bundle(s) -> R2")


# ---------------------------------------------------------------------------
# Bulk search
# ---------------------------------------------------------------------------

def _mmddyyyy(d: date) -> str:
    return d.strftime("%m/%d/%Y")


def search_page(
    session: requests.Session,
    posted_from: date,
    posted_to: date,
    offset: int,
    limit: int = 1000,
) -> tuple[list[dict], int]:
    """One bulk search call. Returns (opportunitiesData, totalRecords)."""
    r = session.get(
        OPP_SEARCH_API,
        params={
            "api_key":    API_KEY,
            "postedFrom": _mmddyyyy(posted_from),
            "postedTo":   _mmddyyyy(posted_to),
            "limit":      limit,
            "offset":     offset,
        },
        timeout=120,
    )
    if r.status_code == 429:
        raise SystemExit("SAM 429 — quota exhausted. Resumes midnight UTC.")
    if r.status_code != 200:
        raise SystemExit(f"SAM HTTP {r.status_code}: {r.text[:300]}")
    data = r.json()
    return data.get("opportunitiesData") or [], int(data.get("totalRecords") or 0)


def iter_opps_in_window(
    session: requests.Session,
    posted_from: date,
    posted_to: date,
    max_calls: int,
    start_offset: int = 0,
) -> Iterator[tuple[dict, int]]:
    """Yield each opp + running page number. Stops at max_calls or when drained."""
    offset = start_offset
    page = 0
    total = None
    while page < max_calls:
        page += 1
        opps, total_records = search_page(session, posted_from, posted_to, offset)
        if total is None:
            total = total_records
            print(f"  window has {total:,} opps total; up to {max_calls} page(s) * 1000 this run")
        if not opps:
            break
        for opp in opps:
            yield opp, page
        if len(opps) < 1000:
            break
        offset += 1000
        time.sleep(1.0)  # light throttle between pages


# ---------------------------------------------------------------------------
# Attachment download + text extraction
# ---------------------------------------------------------------------------

def _extract_pdf_text(pdf_bytes: bytes) -> tuple[int | None, str | None]:
    try:
        reader = PdfReader(io.BytesIO(pdf_bytes))
        parts = []
        for p in reader.pages:
            try:
                parts.append(p.extract_text() or "")
            except Exception:
                parts.append("")
        return len(reader.pages), "\n\n".join(parts).strip()
    except Exception as exc:
        return None, f"[pypdf extract failed: {exc}]"


def _extract_docx_text(docx_bytes: bytes) -> tuple[int | None, str | None]:
    if Document is None:
        return None, "[python-docx not installed]"
    try:
        doc = Document(io.BytesIO(docx_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text]
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text for c in row.cells if c.text]
                if cells:
                    paragraphs.append(" | ".join(cells))
        return len(doc.paragraphs), "\n\n".join(paragraphs).strip()
    except Exception as exc:
        return None, f"[python-docx extract failed: {exc}]"


def _extract_xlsx_text(xlsx_bytes: bytes) -> tuple[int | None, str | None]:
    if load_workbook is None:
        return None, "[openpyxl not installed]"
    try:
        wb = load_workbook(io.BytesIO(xlsx_bytes), read_only=True, data_only=True)
        out = []
        total_rows = 0
        for ws in wb.worksheets:
            out.append(f"### Sheet: {ws.title}")
            for row in ws.iter_rows(values_only=True):
                cells = ["" if v is None else str(v).strip() for v in row]
                if any(cells):
                    out.append(" | ".join(cells))
                    total_rows += 1
            out.append("")
        return total_rows, "\n".join(out).strip()
    except Exception as exc:
        return None, f"[openpyxl extract failed: {exc}]"


def _extract_by_ext(filename: str, content_type: str, data: bytes) -> tuple[int | None, str | None]:
    lower = filename.lower()
    ct = (content_type or "").lower()
    if lower.endswith(".pdf") or ct.startswith("application/pdf"):
        return _extract_pdf_text(data)
    if lower.endswith(".docx") or "officedocument.wordprocessingml" in ct:
        return _extract_docx_text(data)
    if lower.endswith(".xlsx") or "officedocument.spreadsheetml" in ct:
        return _extract_xlsx_text(data)
    return None, None


def download_and_extract(session: requests.Session, opp: dict) -> list[dict]:
    links = opp.get("resourceLinks") or []
    out: list[dict] = []
    for i, url in enumerate(links):
        fetch_url = url
        if "sam.gov" in url and "api_key=" not in url:
            fetch_url = f"{url}{'&' if '?' in url else '?'}api_key={API_KEY}"
        try:
            resp = session.get(fetch_url, timeout=180, stream=True, allow_redirects=True)
        except requests.RequestException as exc:
            out.append({"index": i, "url": url, "error": str(exc)})
            continue
        if resp.status_code != 200:
            out.append({"index": i, "url": url, "error": f"HTTP {resp.status_code}"})
            continue

        data = resp.content
        filename = None
        cd = resp.headers.get("content-disposition", "")
        if "filename=" in cd:
            filename = cd.split("filename=", 1)[1].strip().strip('"').strip("'")
        if not filename:
            filename = re.split(r"[?#]", url.rstrip("/").split("/")[-1])[0] or f"attachment_{i}"

        sha = hashlib.sha256(data).hexdigest()
        pages, text = _extract_by_ext(filename, resp.headers.get("content-type", ""), data)
        chars = len(text) if text else 0

        out.append({
            "index":    i,
            "url":      url,
            "filename": filename,
            "bytes":    len(data),
            "sha256":   sha,
            "pages":    pages,
            "chars":    chars,
            "text":     text,
        })
    return out


# ---------------------------------------------------------------------------
# Labels (regex-only for now — kept minimal on purpose)
# ---------------------------------------------------------------------------

_RE_RTM   = re.compile(r"\brequirements?\s+traceability\s+matrix\b", re.IGNORECASE)
_RE_SHALL = re.compile(r"\bshall\b", re.IGNORECASE)
_RE_AGILE = re.compile(
    r"\b(sprint|agile|scrum|kanban|iteration|backlog|user\s+stor(y|ies)|mvp|working\s+software|ceremon(y|ies)|stand[- ]?up|retrospective)\b",
    re.IGNORECASE,
)
_RE_USER  = re.compile(
    r"\b(end[- ]?users?|stakeholders?|user\s+research|user\s+needs?|user\s+experience|ux)\b",
    re.IGNORECASE,
)


def classify_bundle_text(text: str) -> dict:
    return {
        "mentions_rtm":     bool(_RE_RTM.search(text)),
        "shall_count":      len(_RE_SHALL.findall(text)),
        "has_agile_vocab":  bool(_RE_AGILE.search(text)),
        "has_user_vocab":   bool(_RE_USER.search(text)),
    }


# ---------------------------------------------------------------------------
# Bundle builder
# ---------------------------------------------------------------------------

def build_bundle(opp: dict, attachments: list[dict]) -> dict:
    combined = "\n\n".join((a.get("text") or "") for a in attachments if a.get("text"))
    desc = opp.get("description") or ""
    labels = classify_bundle_text(combined + "\n\n" + desc)

    return {
        "notice_id":          opp.get("noticeId"),
        "fetched_at":         datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%MZ"),
        "labels":             labels,
        "metadata": {
            "title":                 opp.get("title"),
            "type":                  opp.get("type"),
            "base_type":             opp.get("baseType"),
            "naics_code":            opp.get("naicsCode"),
            "classification":        opp.get("classificationCode"),
            "active":                opp.get("active"),
            "solicitation_number":   opp.get("solicitationNumber"),
            "posted_date":           opp.get("postedDate"),
            "response_deadline":     opp.get("responseDeadLine"),
            "archive_date":          opp.get("archiveDate"),
            "set_aside":             opp.get("typeOfSetAside"),
            "set_aside_desc":        opp.get("typeOfSetAsideDescription"),
            "description":           opp.get("description"),
            "full_path_name":        opp.get("fullParentPathName"),
            "department":            (opp.get("fullParentPathName") or "").split(".")[0] if opp.get("fullParentPathName") else None,
            "point_of_contact":      opp.get("pointOfContact"),
            "award":                 opp.get("award"),
            "place_of_performance":  opp.get("placeOfPerformance"),
            "ui_link":               opp.get("uiLink"),
        },
        "attachments":        attachments,
    }


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def _naics_matches(naics: str, prefixes: tuple[str, ...]) -> bool:
    return any(naics.startswith(p) for p in prefixes)


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument("--start-date",
                    help="First-run window start (YYYY-MM-DD). Ignored if state has last_fetched_date.")
    ap.add_argument("--force-window-start",
                    help="Override any existing state and start window from this date (YYYY-MM-DD). Use for backfills.")
    ap.add_argument("--max-api-calls", type=int, default=50,
                    help="Cap on search pages this run (default: 50 = up to 50,000 opps scanned)")
    ap.add_argument("--naics-prefix", nargs="+", default=list(DEFAULT_NAICS_PREFIXES))
    ap.add_argument("--types", nargs="+", default=sorted(DEFAULT_NOTICE_TYPES))
    ap.add_argument("--dry-run", action="store_true",
                    help="List window + first page of matches; no downloads")
    ap.add_argument("--summary-file", default="")
    args = ap.parse_args()

    if not args.dry_run and not API_KEY:
        raise SystemExit("SAM_API_KEY not set")

    STATE_DIR.mkdir(parents=True, exist_ok=True)
    BUNDLE_DIR.mkdir(parents=True, exist_ok=True)
    _maybe_hydrate_from_r2()

    processed: set[str] = set(_load_json(PROCESSED_JSON, []))
    last_fetched = _load_json(LAST_DATE_JSON, None)
    scan_cursor = _load_json(SCAN_CURSOR_JSON, None)
    today = date.today()

    # --force-window-start overrides everything: clears cursor + last_fetched
    # so we scan from the given date. Use for backfills.
    if args.force_window_start:
        posted_from = datetime.strptime(args.force_window_start, "%Y-%m-%d").date()
        posted_to = today
        start_offset = 0
        scan_cursor = None
        if SCAN_CURSOR_JSON.exists():
            SCAN_CURSOR_JSON.unlink()
        if LAST_DATE_JSON.exists():
            LAST_DATE_JSON.unlink()
        print(f"Force-window-start: overriding state, scanning from {posted_from}")
    # scan_cursor wins when present — means a previous run bailed mid-drain
    # (SAM 429 or page cap). We pin posted_from/posted_to from that run so
    # offsets stay stable across days, and resume from the saved offset.
    elif scan_cursor:
        posted_from = datetime.strptime(scan_cursor["posted_from"], "%Y-%m-%d").date()
        posted_to   = datetime.strptime(scan_cursor["posted_to"],   "%Y-%m-%d").date()
        start_offset = int(scan_cursor["offset"])
    elif last_fetched:
        # Overlap by 1 day to catch late-posted amendments.
        posted_from = datetime.strptime(last_fetched, "%Y-%m-%d").date() - timedelta(days=1)
        posted_to = today
        start_offset = 0
    elif args.start_date:
        posted_from = datetime.strptime(args.start_date, "%Y-%m-%d").date()
        posted_to = today
        start_offset = 0
    else:
        posted_from = today - timedelta(days=DEFAULT_LOOKBACK_DAYS)
        posted_to = today
        start_offset = 0

    notice_types = set(args.types)
    naics_prefixes = tuple(args.naics_prefix)

    print(f"Window: {posted_from} → {posted_to}")
    if start_offset:
        print(f"Resuming mid-drain at offset {start_offset:,}")
    print(f"State: {len(processed):,} processed so far")
    print(f"Notice types: {len(notice_types)} kept; NAICS: {list(naics_prefixes)}")

    if args.dry_run:
        session = requests.Session()
        opps, total = search_page(session, posted_from, posted_to, start_offset)
        match = [o for o in opps if _naics_matches((o.get("naicsCode") or ""), naics_prefixes)
                 and o.get("type") in notice_types]
        print(f"\n--dry-run: first page = {len(opps)} opps of {total:,} in window; "
              f"{len(match)} match NAICS+type; {sum(1 for o in match if o.get('noticeId') not in processed)} unseen.")
        return

    session = requests.Session()
    bundles_written: list[str] = []
    pages = 0
    scanned = kept_new = skipped_processed = skipped_type = skipped_naics = 0
    sam_error: str | None = None

    try:
        for opp, page_no in iter_opps_in_window(session, posted_from, posted_to, args.max_api_calls, start_offset=start_offset):
            scanned += 1
            pages = page_no
            nid = opp.get("noticeId")
            if not nid:
                continue

            if nid in processed:
                skipped_processed += 1
                continue
            if opp.get("type") not in notice_types:
                skipped_type += 1
                continue
            naics = opp.get("naicsCode") or ""
            if not _naics_matches(naics, naics_prefixes):
                skipped_naics += 1
                continue

            attachments = download_and_extract(session, opp)
            bundle = build_bundle(opp, attachments)
            (BUNDLE_DIR / f"{nid}.json").write_text(json.dumps(bundle, indent=2, default=str))
            processed.add(nid)
            bundles_written.append(nid)
            kept_new += 1

            if kept_new % 25 == 0:
                print(f"  {kept_new} bundled (page {page_no}, scanned {scanned})")
    except SystemExit as e:
        # 429 or other SAM error — remember it, save what we have, fail at the end.
        sam_error = str(e)
        print(f"{sam_error}")

    # Cursor logic:
    #   - full drain (short page, no error): advance last_fetched_date to
    #     the pinned posted_to; clear scan_cursor.
    #   - otherwise (SAM 429 or page cap): keep last_fetched_date alone,
    #     save scan_cursor with the pinned window + resume offset so the
    #     next run picks up where we stopped instead of re-scanning the
    #     same top-of-window pages.
    full_drain = (sam_error is None
                  and pages > 0
                  and pages < args.max_api_calls)  # exited via len < 1000
    final_offset = start_offset + pages * 1000
    if full_drain:
        _save_json(LAST_DATE_JSON, posted_to.isoformat())
        _clear_scan_cursor()
    elif pages > 0:
        _save_json(SCAN_CURSOR_JSON, {
            "posted_from": posted_from.isoformat(),
            "posted_to":   posted_to.isoformat(),
            "offset":      final_offset,
            "updated_at":  datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%MZ"),
            "reason":      sam_error or "page cap reached",
        })
        print(f"  scan_cursor saved: offset={final_offset:,} (window pinned {posted_from} → {posted_to})")

    _save_json(PROCESSED_JSON, sorted(processed))
    _save_json(QUOTA_JSON, {
        "run_at":              datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%MZ"),
        "window":              f"{posted_from} to {posted_to}",
        "api_pages":           pages,
        "start_offset":        start_offset,
        "end_offset":          final_offset,
        "opps_scanned":        scanned,
        "bundles_written":     len(bundles_written),
        "skipped_already_processed": skipped_processed,
        "skipped_wrong_type":  skipped_type,
        "skipped_wrong_naics": skipped_naics,
        "window_fully_drained": full_drain,
    })

    _push_state_to_r2()
    _push_bundles_to_r2(bundles_written)

    print(f"\nRun complete: {pages} API call(s), {scanned:,} scanned, "
          f"{kept_new} bundled. {len(processed):,} total processed.")

    if args.summary_file:
        lines = [f"## RFP text pipeline (bulk-search)\n"]
        lines.append(f"- Window: **{posted_from} → {posted_to}**")
        lines.append(f"- Offset: **{start_offset:,} → {final_offset:,}**")
        lines.append(f"- API pages spent: **{pages}** / {args.max_api_calls}")
        lines.append(f"- Opps scanned: **{scanned:,}**")
        lines.append(f"- Bundles written: **{len(bundles_written)}**")
        lines.append(f"- Skipped (already processed / wrong type / wrong NAICS): "
                     f"{skipped_processed:,} / {skipped_type:,} / {skipped_naics:,}")
        lines.append(f"- Total processed to date: **{len(processed):,}**")
        lines.append(f"- Window fully drained: **{full_drain}**\n")
        with open(args.summary_file, "a") as f:
            f.write("\n".join(lines) + "\n")

    if sam_error:
        raise SystemExit(f"FAIL: SAM error during run — {sam_error}")
    if kept_new == 0:
        raise SystemExit("FAIL: 0 new bundles written this run.")


if __name__ == "__main__":
    main()
